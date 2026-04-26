from __future__ import annotations

import csv
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = ("txt", "md", "pdf", "docx", "csv", "xlsx")


def supported_upload_types() -> list[str]:
    return list(SUPPORTED_EXTENSIONS)


def read_uploaded_file(file: Any) -> str:
    name = getattr(file, "name", "")
    return read_file_bytes(name, file.getvalue())


def read_file_path(path: str | Path) -> str:
    file_path = Path(path)
    return read_file_bytes(file_path.name, file_path.read_bytes())


def read_file_bytes(name: str, data: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix in {".txt", ".md"}:
        return decode_text(data).strip()
    if suffix == ".pdf":
        return read_pdf(BytesIO(data))
    if suffix == ".docx":
        return read_docx(data)
    if suffix == ".csv":
        return read_csv(data)
    if suffix == ".xlsx":
        return read_xlsx(data)
    return ""


def read_pdf(file: Any) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(file)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    except Exception:
        return ""


def read_docx(data: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(data))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def read_csv(data: bytes) -> str:
    text = decode_text(data)
    reader = csv.reader(StringIO(text))
    return flatten_rows(reader, source_name="CSV")


def read_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
        chunks = []
        for sheet in workbook.worksheets:
            rows = ([clean_cell(cell) for cell in row] for row in sheet.iter_rows(values_only=True))
            sheet_text = flatten_rows(rows, source_name=sheet.title)
            if sheet_text:
                chunks.append(sheet_text)
        workbook.close()
        return "\n\n".join(chunks).strip()
    except Exception:
        return ""


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def clean_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\n", " ").strip()


def flatten_rows(rows: Any, source_name: str, max_rows: int = 250) -> str:
    cleaned_rows = [[clean_cell(cell) for cell in row] for row in rows]
    cleaned_rows = [[cell for cell in row] for row in cleaned_rows if any(cell for cell in row)]
    if not cleaned_rows:
        return ""

    headers = cleaned_rows[0]
    body = cleaned_rows[1:] if len(cleaned_rows) > 1 else cleaned_rows
    chunks = [f"{source_name} structured input"]
    for index, row in enumerate(body[:max_rows], start=1):
        pairs = []
        for cell_index, cell in enumerate(row):
            if not cell:
                continue
            header = headers[cell_index] if cell_index < len(headers) and headers[cell_index] else f"Column {cell_index + 1}"
            pairs.append(f"{header}: {cell}")
        if pairs:
            chunks.append(f"Row {index}: " + "; ".join(pairs))
    if len(body) > max_rows:
        chunks.append(f"Rows omitted for speed: {len(body) - max_rows}")
    return "\n".join(chunks).strip()
